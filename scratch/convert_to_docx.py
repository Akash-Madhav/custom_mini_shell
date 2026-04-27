from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_txt_to_docx(txt_file, docx_file):
    doc = Document()
    
    # Set default style to Courier New for code
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Courier New'
    font.size = Pt(10)

    try:
        with open(txt_file, 'r', encoding='utf-8') as f:
            for line in f:
                # Check for file headers to make them bold or different
                if line.startswith("File: ") or line.startswith("---") or line.startswith("==="):
                    p = doc.add_paragraph()
                    run = p.add_run(line.strip())
                    run.bold = True
                    if line.startswith("File: "):
                        run.font.size = Pt(12)
                        run.font.underline = True
                else:
                    doc.add_paragraph(line.replace('\n', '').replace('\r', ''))

        doc.save(docx_file)
        print(f"Successfully converted {txt_file} to {docx_file}")
    except Exception as e:
        print(f"Error during conversion: {e}")

if __name__ == "__main__":
    convert_txt_to_docx('all_project_code.txt', 'all_project_code.docx')
